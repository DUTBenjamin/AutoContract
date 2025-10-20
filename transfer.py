import json
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn  # 导入命名空间工具

# 检查是否提供了输入和输出文件路径
if len(sys.argv) != 3:
    print("Usage: python script.py <input_json_file> <output_docx_file>")
    sys.exit(1)

input_file = sys.argv[1]
output_file = sys.argv[2]

# 加载 JSON 数据
with open(input_file, "r", encoding="utf-8") as file:
    data = json.load(file)

# 定义章节标题映射及中文翻译，添加新的映射关系
section_titles = [
    ("<Definitions>", "第一条  名词和术语"),
    ("<LicenseGrant>", "第二条  许可的授予"),
    ("<PaymentTerms>", "第三条  许可费及支付方式"),
    ("<TechnicalDelivery>", "第四条  技术资料的交付与验收"),
    ("<TechnicalService>", "第五条  技术服务与培训"),
    ("<Confidentiality>", "第六条  保密条款"),
    ("<Improvement>", "第七条  后续改进成果的提供与分享"),
    ("<Representations>", "第八条  陈述与保证"),
    ("<TechnologyExport>", "第九条  技术进出口"),
    ("<InfringementResponse>", "第十条  侵权应对及共同维权"),
    ("<PatentInvalidity>", "第十一条  专利权被宣告无效（或专利申请被驳回）的处理"),
    ("<ForceMajeure>", "第十二条  不可抗力"),
    ("<Delivery>", "第十三条  送达"),
    ("<Breach>", "第十四条  违约与损害赔偿"),
    ("<Tax>", "第十五条  税费"),
    ("<DisputeResolution>", "第十六条  争议解决"),
    ("<ContractEffectiveness>", "第十七条 合同的生效、变更与终止"),
    ("<ServiceContent>", "一、技术服务内容"),
    ("<TrainingArrangement>", "二、技术培训安排"),
    ("<LicensorRepresentations>", "许可方的陈述与保证"),
    ("<LicenseeRepresentations>", "被许可方的陈述与保证"),
    ("<Representation>", ""),
    ("<ComplianceStatement>", ""),
    ("<RightsAndObligations>",""),
    ("<Description>",""),

]


# 清理不必要的 Markdown 语法和符号
def clean_extra_symbols(text):
    # 移除 Markdown 加粗和斜体符号
    text = re.sub(r'\*\*+(.*?)\*\*+', r'\1', text)
    text = re.sub(r'\*+(.*?)\*+', r'\1', text)

    # 移除项目符号或其他不必要的符号
    text = re.sub(r'^\s*[-\*\+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\s{2,}', ' ', text)

    # 移除剩余的内联代码
    text = re.sub(r'`(.*?)`', r'\1', text)

    return text


# 添加加粗样式并设置字体和字号
def add_bold(paragraph, text, font_size=Pt(16)):
    """Add bold to a given paragraph with specified font size and font name."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = font_size
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')


# 创建 Word 文档
doc = Document()

# 设置默认字体为仿宋，字号为 12
style = doc.styles['Normal']
font = style.font
font.name = '仿宋'
font.size = Pt(12)

# 检查数据是否是字典类型，如果是，转换为列表
if isinstance(data, dict):
    data = [data]

# 遍历章节标题并提取每个章节的内容
for section_tag, section_name in section_titles:
    for entry in data:
        # 尝试获取内容，优先使用 ground_truth，然后是 model_output
        content = ""
        if isinstance(entry, dict):
            content = entry.get("ground_truth", entry.get("model_output", ""))
        elif isinstance(entry, str):
            content = entry

        if section_tag in content:
            # 添加章节标题（替换为中文标题），设置为加粗，字号 16
            title_paragraph = doc.add_paragraph()
            add_bold(title_paragraph, section_name, Pt(16))

            # 提取章节标签之间的内容
            start_idx = content.find(section_tag) + len(section_tag)
            end_idx = content.find(f"</{section_tag.strip('<>')}>")
            section_content = content[start_idx:end_idx].strip()

            # 替换英文标签为中文标题
            for tag, title in section_titles:
                open_tag = tag
                close_tag = f"</{tag.strip('<>')}>"
                section_content = section_content.replace(open_tag, title)
                section_content = section_content.replace(close_tag, "")

            # 使用正则表达式按连续换行符分割内容
            paragraphs = re.split(r'\n+', section_content)

            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue

                # 清理当前段落的符号
                cleaned_paragraph = clean_extra_symbols(paragraph)

                # 检查是否为编号列表
                numbered_list_match = re.match(r'^(\d+\.)\s+(.*)$', cleaned_paragraph)
                if numbered_list_match:
                    num = numbered_list_match.group(1)
                    content_part = numbered_list_match.group(2)
                    para = doc.add_paragraph()
                    para_run_num = para.add_run(num)
                    para_run_num.font.name = '仿宋'
                    para_run_num._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    para_run_content = para.add_run(" " + content_part)
                    para_run_content.font.name = '仿宋'
                    para_run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                else:
                    para = doc.add_paragraph()
                    para_run = para.add_run(cleaned_paragraph)
                    para_run.font.name = '仿宋'
                    para_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

                # 在每个段落之后添加空行以实现换行
                doc.add_paragraph()

# 保存 Word 文档
doc.save(output_file)

print(output_file)